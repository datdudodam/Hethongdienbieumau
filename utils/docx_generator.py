
from docx import Document
import os
import uuid
from urllib.parse import quote
from config.config import UPLOADS_DIR
import datetime
from models.data_model import load_form_history, save_form_history

def generate_docx(form_data, doc_path, custom_filename=None):
    """
    Tạo tài liệu docx từ mẫu và dữ liệu form
    """
    # Validate template document
    if not doc_path:
        return {"error": "Không tìm thấy tài liệu mẫu. Vui lòng tải lên tài liệu trước."}, 404
    if not os.path.exists(doc_path):
        return {"error": "Tài liệu mẫu không còn tồn tại. Vui lòng tải lên lại tài liệu."}, 404
    if not os.path.isfile(doc_path) or not doc_path.endswith('.docx'):
        return {"error": "Tài liệu mẫu không hợp lệ. Vui lòng tải lên tài liệu DOCX."}, 400

    # Process document
    try:
        doc = Document(doc_path)
    except Exception as e:
        print(f"Error loading document template: {str(e)}")
        return {"error": "Không thể mở tài liệu mẫu. Tài liệu có thể bị hỏng hoặc không đúng định dạng."}, 500

    # Replace placeholders
    try:
        for paragraph in doc.paragraphs:
            original_text = paragraph.text
            modified_text = original_text
            for field_code, value in form_data.items():
                if field_code in modified_text:
                    try:
                        # Đảm bảo giá trị là chuỗi và xử lý None
                        safe_value = str(value) if value is not None else ''
                        modified_text = modified_text.replace(field_code, safe_value)
                    except Exception as e:
                        print(f"Error replacing field {field_code}: {str(e)}")
                        return {"error": f"Lỗi khi thay thế dữ liệu cho trường {field_code}"}, 500
            if modified_text != original_text:
                paragraph.text = modified_text

        # Xử lý các bảng nếu có
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        original_text = paragraph.text
                        modified_text = original_text
                        for field_code, value in form_data.items():
                            if field_code in modified_text:
                                try:
                                    safe_value = str(value) if value is not None else ''
                                    modified_text = modified_text.replace(field_code, safe_value)
                                except Exception as e:
                                    print(f"Error replacing field in table {field_code}: {str(e)}")
                                    return {"error": f"Lỗi khi thay thế dữ liệu trong bảng cho trường {field_code}"}, 500
                        if modified_text != original_text:
                            paragraph.text = modified_text
    except Exception as e:
        print(f"Error replacing placeholders: {str(e)}")
        return {"error": "Lỗi khi thay thế dữ liệu trong tài liệu"}, 500

    # Generate output document
    temp_filename = f'temp_output_{uuid.uuid4()}.docx'
    temp_doc_path = os.path.join(UPLOADS_DIR, temp_filename)

    try:
        # Ensure uploads directory exists
        if not os.path.exists(UPLOADS_DIR):
            os.makedirs(UPLOADS_DIR)

        # Save document
        doc.save(temp_doc_path)
        with open(temp_doc_path, 'rb') as f:
            doc_data = f.read()

        # Prepare download filename
        download_filename = custom_filename if custom_filename else os.path.basename(doc_path)
        if not download_filename.endswith('.docx'):
            download_filename += '.docx'
        
        # Xử lý ký tự đặc biệt trong tên file
        ascii_filename = download_filename.encode('ascii', 'ignore').decode()
        utf8_filename = quote(download_filename.encode('utf-8'))
        
        # Trả về thông tin để tạo response
        return {
            "success": True,
            "doc_data": doc_data,
            "ascii_filename": ascii_filename,
            "utf8_filename": utf8_filename,
            "temp_doc_path": temp_doc_path
        }, 200

    except Exception as e:
        print(f"Error processing document: {str(e)}")
        return {"error": "Lỗi khi xử lý tài liệu. Vui lòng thử lại."}, 500