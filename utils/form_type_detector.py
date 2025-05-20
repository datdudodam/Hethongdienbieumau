import os
import re
from docx import Document
import logging

logger = logging.getLogger(__name__)

class FormTypeDetector:
    """
    Lớp phát hiện loại biểu mẫu dựa trên nội dung và tiêu đề
    """
    def __init__(self):
        # Từ điển ánh xạ các từ khóa với loại biểu mẫu
        self.form_type_keywords = {
            'báo cáo tuần': 'Báo cáo tuần',
            'báo cáo tháng': 'Báo cáo tháng',
            'báo cáo quý': 'Báo cáo quý',
            'báo cáo năm': 'Báo cáo năm',
            'đơn xin việc': 'Đơn xin việc',
            'đơn xin nghỉ': 'Đơn xin nghỉ phép',
            'đơn xin phép': 'Đơn xin nghỉ phép',
            'hợp đồng': 'Hợp đồng',
            'biên bản': 'Biên bản',
            'kế hoạch': 'Kế hoạch',
            'tờ trình': 'Tờ trình',
            'công văn': 'Công văn',
            'giấy mời': 'Giấy mời',
            'thông báo': 'Thông báo',
            'quyết định': 'Quyết định',
        }
    
    def detect_form_type(self, doc_path):
        """
        Phát hiện loại biểu mẫu dựa trên nội dung tài liệu
        """
        try:
            # Kiểm tra đường dẫn tài liệu
            if not os.path.exists(doc_path):
                logger.error(f"Không tìm thấy tài liệu: {doc_path}")
                return "Unknown"
            
            # Lấy tên file để phân tích
            file_name = os.path.basename(doc_path).lower()
            
            # Kiểm tra từ khóa trong tên file
            for keyword, form_type in self.form_type_keywords.items():
                if keyword in file_name:
                    return form_type
            
            # Nếu không tìm thấy từ tên file, phân tích nội dung
            try:
                doc = Document(doc_path)
                
                # Lấy văn bản từ tài liệu
                full_text = ""
                for para in doc.paragraphs:
                    full_text += para.text.lower() + " "
                
                # Kiểm tra từ khóa trong nội dung
                for keyword, form_type in self.form_type_keywords.items():
                    if keyword in full_text:
                        return form_type
                
                # Nếu có tiêu đề, sử dụng tiêu đề làm form_type
                if doc.paragraphs and doc.paragraphs[0].text.strip():
                    title = doc.paragraphs[0].text.strip()
                    # Nếu tiêu đề quá dài, rút gọn
                    if len(title) > 50:
                        title = title[:50] + "..."
                    return title
                
            except Exception as e:
                logger.error(f"Lỗi khi phân tích nội dung tài liệu: {str(e)}")
            
            # Nếu không thể xác định, sử dụng tên file làm form_type
            return os.path.splitext(file_name)[0]
            
        except Exception as e:
            logger.error(f"Lỗi khi phát hiện loại biểu mẫu: {str(e)}")
            return "Unknown"