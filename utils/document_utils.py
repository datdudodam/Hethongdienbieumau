import re
from docx import Document
import os
import uuid
from werkzeug.utils import secure_filename
from config.config import UPLOADS_DIR
from flask import session
from transformers import AutoTokenizer, AutoModelForTokenClassification
from transformers import pipeline
import spacy

# Load models
model_name = "Davlan/bert-base-multilingual-cased-ner-hrl"
tokenizer = AutoTokenizer.from_pretrained(model_name)
model = AutoModelForTokenClassification.from_pretrained(model_name)
ner_pipeline = pipeline("ner", model=model, tokenizer=tokenizer, aggregation_strategy="simple")

# Load spaCy model for Vietnamese (nếu có) hoặc English
try:
    nlp = spacy.load("vi_core_news_lg")  # Hoặc "en_core_web_sm" nếu không có tiếng Việt
except:
    nlp = spacy.load("en_core_web_sm")

# Biến toàn cục để lưu đường dẫn tài liệu hiện tại
doc_path = None

def load_document(doc_path):
    """Tải nội dung từ tài liệu docx"""
    doc = Document(doc_path)
    return "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])

def clean_label(text):
    """Làm sạch nhãn để tránh nhiễu đầu ra"""
    text = re.sub(r"\[_\d+_\]", "", text)
    text = re.sub(r"[^\w\sÀ-ỹ]", "", text)  # Giữ lại chữ cái tiếng Việt
    text = re.sub(r"\s+", " ", text).strip()
    return text

def extract_key_nouns(text):
    """
    Trích xuất danh từ chính từ văn bản sử dụng spaCy
    """
    doc = nlp(text)
    nouns = []
    
    # Ưu tiên các danh từ đơn và cụm danh từ
    for chunk in doc.noun_chunks:
        # Loại bỏ các từ không quan trọng (mạo từ, giới từ, v.v.)
        clean_chunk = " ".join([token.text for token in chunk 
                              if token.pos_ in ["NOUN", "PROPN", "ADJ"]])
        if clean_chunk:
            nouns.append(clean_chunk)
    
    # Nếu không tìm thấy cụm danh từ, lấy các danh từ đơn lẻ
    if not nouns:
        nouns = [token.text for token in doc if token.pos_ in ["NOUN", "PROPN"]]
    
    return nouns

def determine_field_name(text: str, window_size: int = 50) -> str:
    """Xác định tên trường từ văn bản sử dụng heuristic và AI"""
    cleaned_text = clean_label(text)
    if not cleaned_text:
        return ""
    
      # 1. Xóa số thứ tự ở đầu: vd. "1. ", "2 ", "3-" → bỏ
    cleaned_text = re.sub(r"^\s*\d+[\.\-\)]?\s*", "", cleaned_text).strip()
    words = cleaned_text.split()
     # 2. Nếu có số thứ tự ở cuối → giữ nguyên nhãn (trường hợp như "Tên nhân viên 1")
    if re.search(r"\d+$", cleaned_text):
        return cleaned_text.strip()
    # 3. Nếu nhãn ngắn → giữ nguyên
    if len(words) <= 5:
        return cleaned_text.strip()
    # 3. Trích xuất danh từ chính bằng spaCy
    key_nouns = extract_key_nouns(cleaned_text)
    if key_nouns:
        # Lấy danh từ cuối cùng (thường là danh từ chính)
        main_noun = key_nouns[-1]
        
        # Nếu có nhiều hơn 1 danh từ, kết hợp với danh từ trước đó nếu ngắn
        if len(key_nouns) > 1 and len(main_noun.split()) < 3:
            combined = f"{key_nouns[-2]} {main_noun}"
            if len(combined.split()) <= 4:
                return combined
        
        return main_noun
    
    # 4. Sử dụng NER nếu không trích xuất được danh từ
    ner_results = ner_pipeline(cleaned_text[:window_size])
    important_entities = [
        entity["word"].strip() for entity in ner_results 
        if entity["entity_group"] in {"PER", "ORG", "LOC", "MISC"}
    ]
    
    if important_entities:
        return " ".join(important_entities[-2:]).strip()  # Lấy tối đa 2 thực thể cuối
    
    # 5. Fallback: Lấy 3-4 từ cuối cùng nếu các phương pháp trên thất bại
    
    return " ".join(words[:3]).strip()







def extract_fields(paragraphs, window_size=50):
    """Trích xuất tên trường từ các đoạn văn bản"""
    field_pattern = r"\[_\d+_\]|_{4,}|\.{4,}|\[fill\]|\[\d+\]"
    fields = []

    for para in paragraphs:
        text = para.text.strip()
        if not text:
            continue

        matches = list(re.finditer(field_pattern, text))
        prev_match_end = 0

        for match in matches:
            field_code = match.group()
            match_start = match.start()
            match_end = match.end()

            # Lấy đoạn trước mã trường (context)
            raw_context = text[prev_match_end:match_start].strip()
            prev_match_end = match_end
            if not raw_context:
                para_idx = para._element.getparent().index(para._element)
                if para_idx > 0:
                    prev_para = para._element.getparent()[para_idx - 1]
                    raw_context = prev_para.text.strip() if prev_para.text else ""
            # Làm sạch context
            cleaned_context = clean_label(raw_context)
            if not cleaned_context:
                continue

            # Xác định tên trường
            field_name = determine_field_name(cleaned_context, window_size)
            if not field_name:
                continue

            fields.append({
                "field_name": field_name,
                "field_code": field_code,
                "source": "text",
                "position": (para._element.getparent().index(para._element), 0),
                "raw_context": raw_context  # Lưu thêm context gốc để debug
            })

    return fields



def extract_fields_from_tables(tables):
    """Trích xuất trường từ các bảng"""
    fields = []
    field_patterns = [r"\[_\d+_\]", r"\.{4,}", r"_{4,}", r"\[fill\]", r"\[\s*\d+\s*\]"]
    
    for table_idx, table in enumerate(tables):
        if not table.rows:
            continue
            
        num_cols = len(table.rows[0].cells)
        
        if num_cols == 2:
            # Bảng 2 cột: cột 1 là nhãn, cột 2 là trường nhập
            for row_idx, row in enumerate(table.rows):
                label_cell = row.cells[0].text.strip()
                if not label_cell:
                    continue
                    
                field_code = None
                for pattern in field_patterns:
                    match = re.search(pattern, row.cells[1].text.strip())
                    if match:
                        field_code = match.group()
                        break
                if not field_code:
                    continue
                    
                field_name = determine_field_name(clean_label(label_cell), window_size=40)
                if field_name:
                    fields.append({
                        "field_name": field_name,
                        "field_code": field_code,
                        "source": "table",
                        "position": (table_idx, row_idx, 1)
                    })
                    
        elif num_cols > 2:
            # Bảng nhiều cột: dòng đầu là header, các dòng sau chứa trường nhập
            header_cells = [cell.text.strip() for cell in table.rows[0].cells]
            
            for row_idx, row in enumerate(table.rows[1:], start=1):
                for col_idx in range(num_cols):
                    cell_text = row.cells[col_idx].text.strip()
                    field_code = None
                    for pattern in field_patterns:
                        match = re.search(pattern, cell_text)
                        if match:
                            field_code = match.group()
                            break
                    if not field_code:
                        continue
                        
                    label = header_cells[col_idx] if col_idx < len(header_cells) else f"Cột {col_idx+1}"
                    field_name = determine_field_name(clean_label(label), window_size=40)
                    if field_name:
                        fields.append({
                            "field_name": field_name,
                            "field_code": field_code,
                            "source": "table",
                            "position": (table_idx, row_idx, col_idx)
                        })
    
    return fields

# [Các hàm còn lại giữ nguyên như extract_all_fields, upload_document, get_doc_path, set_doc_path]

def extract_all_fields(doc_path):
    """
    Trích xuất tất cả các trường từ tài liệu theo đúng thứ tự xuất hiện
    (bao gồm cả văn bản và bảng)
    """
    doc = Document(doc_path)
    
    # Lấy tất cả các phần tử trong tài liệu theo thứ tự
    all_elements = []
    for element in doc.element.body:
        if element.tag.endswith('p'):
            all_elements.append(('paragraph', element))
        elif element.tag.endswith('tbl'):
            all_elements.append(('table', element))
    
    # Trích xuất trường theo thứ tự
    fields = []
    seen_codes = set()
    
    for elem_type, element in all_elements:
        if elem_type == 'paragraph':
            # Tạo paragraph object từ element
            para = next(p for p in doc.paragraphs if p._element is element)
            para_fields = extract_fields([para])
            
            for field in para_fields:
                if field['field_code'] not in seen_codes:
                    fields.append(field)
                    seen_codes.add(field['field_code'])
                    
        elif elem_type == 'table':
            # Tìm table object tương ứng với element
            table = None
            for t in doc.tables:
                if t._element is element:
                    table = t
                    break
                    
            if table:
                table_fields = extract_fields_from_tables([table])
                
                for field in table_fields:
                    if field['field_code'] not in seen_codes:
                        fields.append(field)
                        seen_codes.add(field['field_code'])
    
    return fields

def upload_document(file):
    """
    Xử lý tải lên tài liệu và kiểm tra giới hạn upload của người dùng
    """
    global doc_path
    
    if file.filename == '':
        return {"error": "No selected file"}, 400
    
    if not file.filename.endswith('.docx'):
        return {"error": "Only DOCX files are allowed"}, 400
    
    # Kiểm tra giới hạn upload của người dùng
    from flask_login import current_user
    if current_user.is_authenticated:
        # Kiểm tra loại gói đăng ký
        if current_user.subscription_type == 'free':
            # Kiểm tra số lần upload còn lại
            if current_user.free_downloads_left <= 0:
                return {"error": "Bạn đã sử dụng hết lượt upload miễn phí. Vui lòng nâng cấp lên gói VIP để tiếp tục sử dụng.", "upgrade_required": True}, 403
            
            # Giảm số lần upload còn lại
            from models.user import db
            current_user.free_downloads_left -= 1
            db.session.commit()
        elif current_user.subscription_type == 'standard':
            # Kiểm tra số lần upload trong tháng
            if current_user.monthly_download_count >= 100:
                return {"error": "Bạn đã sử dụng hết 100 lượt upload trong tháng. Vui lòng nâng cấp lên gói VIP để không giới hạn số lần upload.", "upgrade_required": True}, 403
            
            # Tăng số lần upload trong tháng
            from models.user import db
            current_user.monthly_download_count += 1
            db.session.commit()
        # Gói VIP không giới hạn số lần upload
    
    filename = secure_filename(str(uuid.uuid4()) + '_' + file.filename)
    filepath = os.path.join(UPLOADS_DIR, filename)
    file.save(filepath)
    
    # Lưu đường dẫn vào cả session và biến toàn cục
    set_doc_path(filepath)
    
    # Xác định loại biểu mẫu
    from utils.form_type_detector import FormTypeDetector
    detector = FormTypeDetector()
    form_type = detector.detect_form_type(filepath)
    
    # Trả về thông tin về số lần upload còn lại nếu là gói miễn phí
    if current_user.is_authenticated and current_user.subscription_type == 'free':
        return {
            "message": "File uploaded successfully", 
            "filename": filename,
            "free_downloads_left": current_user.free_downloads_left,
            "form_type": form_type
        }, 200
    
    return {"message": "File uploaded successfully", "filename": filename, "form_type": form_type}, 200

def get_doc_path():
    """
    Trả về đường dẫn tài liệu hiện tại từ session hoặc biến toàn cục
    """
    try:
        # Ưu tiên lấy từ session nếu có
        if 'doc_path' in session:
            return session['doc_path']
        return doc_path
    except Exception as e:
        print(f"Error in get_doc_path: {str(e)}")
        return doc_path

def set_doc_path(path):
    """
    Thiết lập đường dẫn tài liệu hiện tại vào session và biến toàn cục
    """
    global doc_path
    try:
        # Lưu vào cả session và biến toàn cục
        session['doc_path'] = path
    except Exception as e:
        print(f"Error setting doc_path in session: {str(e)}")
    # Vẫn lưu vào biến toàn cục để đảm bảo tương thích ngược
    doc_path = path