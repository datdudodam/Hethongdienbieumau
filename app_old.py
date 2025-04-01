import re
import json
from collections import defaultdict, Counter
from docx import Document
from flask import Flask, render_template, request, jsonify
import os
from openai import OpenAI
import datetime

app = Flask(__name__)

# OpenAI Configuration
# Replace 'your-api-key-here' with your actual OpenAI API key
client = OpenAI(api_key='your-api-key-here')

import uuid
from werkzeug.utils import secure_filename

db_path = "form_data.json"
form_history_path = "form_history.json"
uploads_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "uploads")
if not os.path.exists(uploads_dir):
    os.makedirs(uploads_dir)

doc_path = None

def load_document(doc_path):
    doc = Document(doc_path)
    return "\n".join([para.text.strip() for para in doc.paragraphs])

def extract_fields(text):
    matches = re.finditer(r"\[_\d+_\]", text)
    fields = []
    special_words = {"ngày", "tháng", "năm"}
    
    for match in matches:
        end_index = match.start()
        field_name = ""
        i = end_index - 1
        while i >= 0 and text[i] == " ":
            i -= 1
        
        for word in special_words:
            if text[max(0, i - len(word) + 1):end_index].strip().lower() == word:
                field_name = word
                break
        
        if not field_name:
            for j in range(i, -1, -1):
                if text[j].isupper():
                    field_name = text[j:end_index].strip()
                    break
        
        if field_name:
            fields.append({"field_name": field_name, "field_code": match.group()})
    return fields

def load_db():
    if os.path.exists(db_path):
        with open(db_path, "r", encoding="utf-8") as f:
            return json.load(f)
    return []

def save_db(data):
    with open(db_path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=4)

def generate_suggestions(db_data, field_code=None):
    suggestions = defaultdict(Counter)  # Dùng Counter để đếm tần suất các giá trị

    # Tổng hợp dữ liệu từ cơ sở dữ liệu
    for entry in db_data:
        for key, value in entry["data"].items():
            if field_code is None or key == field_code:
                suggestions[key][value] += 1  # Đếm số lần xuất hiện

    result = {}
    for key, counter in suggestions.items():
        # Lấy 3 giá trị phổ biến nhất từ lịch sử nhập liệu
        most_common_values = [val for val, _ in counter.most_common(3)]

        # Gọi OpenAI API để lấy thêm gợi ý thông minh
        gpt_suggestions = []
        if len(counter) > 1:  # Chỉ gọi GPT nếu có đủ dữ liệu để phân tích
            try:
                prompt = (
                    f"Dựa trên các giá trị lịch sử cho trường '{key}': {list(counter.keys())}, "
                    "hãy gợi ý 5 giá trị phù hợp nhất cho trường này. Chỉ trả về danh sách, không kèm theo giải thích."
                )
                response = client.chat.completions.create(
                    model="gpt-4",
                    messages=[
                        {"role": "system", "content": "You are a helpful assistant."},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.6,
                    max_tokens=100
                )

                # Xử lý kết quả GPT trả về
                gpt_suggestions = [s.strip() for s in response.choices[0].message.content.split(",") if s.strip()]
                gpt_suggestions = gpt_suggestions[:5]  # Giới hạn 5 giá trị

            except Exception as e:
                print(f"Lỗi OpenAI API: {e}")

        # Hợp nhất giá trị lịch sử và gợi ý từ GPT, loại bỏ trùng lặp
        combined_suggestions = list(dict.fromkeys(most_common_values + gpt_suggestions))[:5]
        result[key] = combined_suggestions

    return result if field_code is None else result.get(field_code, [])

@app.route('/get_suggestions', methods=['POST'])
def get_suggestions():
    field_code = request.json.get('field_code')
    if not field_code:
        return jsonify({'error': 'Field code is required'}), 400

    db_data = load_db()
    suggestions = generate_suggestions(db_data, field_code)

    return jsonify({'field_code': field_code, 'suggestions': suggestions})

@app.route('/get_gpt_suggestions', methods=['POST'])
def get_gpt_suggestions():
    field_code = request.json.get('field_code')
    if not field_code:
        return jsonify({'error': 'Field code is required'}), 400

    if not client.api_key or client.api_key == 'your-api-key-here':
        return jsonify({'error': 'OpenAI API key is not configured'}), 500

    db_data = load_db()
    suggestions = []

    try:
        prompt = f"Hãy gợi ý 5 giá trị phù hợp nhất cho trường '{field_code}'. Chỉ trả về danh sách, không kèm theo giải thích."
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "Bạn là trợ lý AI giúp đề xuất giá trị phù hợp cho các trường trong biểu mẫu."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=100
        )

        suggestions = [s.strip() for s in response.choices[0].message.content.split(",") if s.strip()]
        suggestions = suggestions[:5]  # Giới hạn 5 giá trị

    except Exception as e:
        print(f"Lỗi xác thực OpenAI API: {e}")
        return jsonify({'error': 'Lỗi xác thực API key'}), 401
    except Exception as e:
        print(f"Lỗi OpenAI API: {e}")
        return jsonify({'error': 'Có lỗi xảy ra khi gọi API'}), 500

    return jsonify({'field_code': field_code, 'suggestions': suggestions})

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    
    if not file.filename.endswith('.docx'):
        return jsonify({'error': 'Only DOCX files are allowed'}), 400
    
    filename = secure_filename(str(uuid.uuid4()) + '_' + file.filename)
    filepath = os.path.join(uploads_dir, filename)
    file.save(filepath)
    
    global doc_path
    doc_path = filepath
    
    return jsonify({'message': 'File uploaded successfully', 'filename': filename})

@app.route('/')
def index():
    return render_template("TrangChu.html")

@app.route('/get-recent-forms')
def get_recent_forms():
    try:
        form_history = load_form_history()
        query = request.args.get('query', '')
        
        # Format the forms for display
        formatted_forms = []
        for form in form_history:
            # Format the date for display
            timestamp = datetime.datetime.fromisoformat(form['timestamp'])
            formatted_date = timestamp.strftime('%d/%m/%Y %H:%M')
            
            # Create a formatted form object
            formatted_form = {
                'id': os.path.basename(form['path']).split('.')[0],
                'name': form['name'],
                'date': formatted_date,
                'path': form['path']
            }
            
            # Add form data if available
            if 'form_data' in form:
                formatted_form['form_data'] = form['form_data']
            
            # Filter by search query if provided
            if query.lower() in form['name'].lower():
                formatted_forms.append(formatted_form)
        
        # If no query, return all forms
        if not query:
            formatted_forms = formatted_forms
            
        # Sort by date (newest first)
        formatted_forms.sort(key=lambda x: x['date'], reverse=True)
        
        return jsonify({'forms': formatted_forms})
    except Exception as e:
        print(f"Error loading recent forms: {str(e)}")
        return jsonify({'error': 'Failed to load recent forms'}), 500

@app.route('/form')
def form():
    if not doc_path:
        return jsonify({'error': 'No document uploaded'}), 400
        
    text = load_document(doc_path)
    fields = extract_fields(text)
    db_data = load_db()
    
    suggestions = generate_suggestions(db_data) if db_data else {}
    
    return render_template("index.html", fields=fields, suggestions=suggestions)

@app.route('/submit', methods=['POST'])
def submit():
    try:
        form_data = request.form.to_dict()
        if not form_data:
            return jsonify({"error": "Không có dữ liệu được gửi"}), 400
            
        db_data = load_db()
        db_data.append({"data": form_data})
        save_db(db_data)
        
        return jsonify({"message": "Dữ liệu đã được gửi thành công!", "data": form_data})
    except Exception as e:
        print(f"Error submitting form data: {str(e)}")
        return jsonify({"error": "Có lỗi xảy ra khi lưu dữ liệu"}), 500

@app.route('/generate-docx', methods=['POST'])
def generate_docx():
    try:
        # Kiểm tra và xử lý dữ liệu từ request
        request_data = None
        form_data = None
        
        if request.is_json:
            try:
                request_data = request.get_json()
                if request_data and isinstance(request_data, dict):
                    form_data = {k: v for k, v in request_data.items() if k != 'filename'}
            except Exception as e:
                print(f"Error parsing JSON: {str(e)}")
                return jsonify({'error': 'Dữ liệu không hợp lệ'}), 400
        else:
            try:
                form_data = request.form.to_dict()
            except Exception as e:
                print(f"Error parsing form data: {str(e)}")
                return jsonify({'error': 'Không thể đọc dữ liệu form'}), 400

        # Nếu không có dữ liệu form, sử dụng dữ liệu mới nhất từ DB
        if not form_data or not any(form_data.values()):
            db_data = load_db()
            if not db_data:
                return jsonify({'error': 'Không có dữ liệu để tạo tài liệu'}), 400
            form_data = db_data[-1]['data']

        custom_filename = request_data.get('filename') if request_data else None

        # Validate template document
        if not doc_path:
            return jsonify({'error': 'Không tìm thấy tài liệu mẫu. Vui lòng tải lên tài liệu trước.'}), 404
        if not os.path.exists(doc_path):
            return jsonify({'error': 'Tài liệu mẫu không còn tồn tại. Vui lòng tải lên lại tài liệu.'}), 404
        if not os.path.isfile(doc_path) or not doc_path.endswith('.docx'):
            return jsonify({'error': 'Tài liệu mẫu không hợp lệ. Vui lòng tải lên tài liệu DOCX.'}), 400

        # Save form history with form data
        try:
            form_history = load_form_history()
            form_info = {
                "name": custom_filename if custom_filename else os.path.basename(doc_path),
                "type": "docx",
                "timestamp": datetime.datetime.now().isoformat(),
                "path": doc_path,
                "form_data": form_data  # Save the form data with the history
            }
            form_history.append(form_info)
            save_form_history(form_history)
        except Exception as e:
            print(f"Error saving form history: {str(e)}")

        # Process document
        try:
            doc = Document(doc_path)
        except Exception as e:
            print(f"Error loading document template: {str(e)}")
            return jsonify({'error': 'Không thể mở tài liệu mẫu. Tài liệu có thể bị hỏng hoặc không đúng định dạng.'}), 500

        # Replace placeholders
        try:
            for paragraph in doc.paragraphs:
                original_text = paragraph.text
                modified_text = original_text
                for field_code, value in form_data.items():
                    if field_code in modified_text:
                        try:
                            # Đảm bảo giá trị là chuỗi và xử lý None
                            safe_value = str(value) if value is not None else ''
                            modified_text = modified_text.replace(field_code, safe_value)
                        except Exception as e:
                            print(f"Error replacing field {field_code}: {str(e)}")
                            return jsonify({'error': f'Lỗi khi thay thế dữ liệu cho trường {field_code}'}), 500
                if modified_text != original_text:
                    paragraph.text = modified_text

            # Xử lý các bảng nếu có
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            original_text = paragraph.text
                            modified_text = original_text
                            for field_code, value in form_data.items():
                                if field_code in modified_text:
                                    try:
                                        safe_value = str(value) if value is not None else ''
                                        modified_text = modified_text.replace(field_code, safe_value)
                                    except Exception as e:
                                        print(f"Error replacing field in table {field_code}: {str(e)}")
                                        return jsonify({'error': f'Lỗi khi thay thế dữ liệu trong bảng cho trường {field_code}'}), 500
                            if modified_text != original_text:
                                paragraph.text = modified_text
        except Exception as e:
            print(f"Error replacing placeholders: {str(e)}")
            return jsonify({'error': 'Lỗi khi thay thế dữ liệu trong tài liệu'}), 500

        # Generate output document
        temp_filename = f'temp_output_{uuid.uuid4()}.docx'
        temp_doc_path = os.path.join(uploads_dir, temp_filename)

        try:
            # Ensure uploads directory exists
            if not os.path.exists(uploads_dir):
                os.makedirs(uploads_dir)

            # Save and process document
            doc.save(temp_doc_path)
            with open(temp_doc_path, 'rb') as f:
                doc_data = f.read()

            # Prepare response
            response = app.make_response(doc_data)
            response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            
            # Xử lý tên file tải xuống
            download_filename = custom_filename if custom_filename else os.path.basename(doc_path)
            if not download_filename.endswith('.docx'):
                download_filename += '.docx'
            
            # Xử lý ký tự đặc biệt trong tên file
            from urllib.parse import quote
            ascii_filename = download_filename.encode('ascii', 'ignore').decode()
            utf8_filename = quote(download_filename.encode('utf-8'))
            
            # Thiết lập headers cho response
            response.headers.set('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response.headers.set('Content-Disposition', f'attachment; filename="{ascii_filename}"; filename*=UTF-8\'\'{utf8_filename}')
            response.headers.set('Content-Length', str(len(doc_data)))
            response.headers.set('Cache-Control', 'no-cache, no-store, must-revalidate')
            response.headers.set('Pragma', 'no-cache')
            response.headers.set('Expires', '0')
            response.headers.set('Access-Control-Expose-Headers', 'Content-Disposition')
            response.headers.set('X-Content-Type-Options', 'nosniff')

            # Return the response object
            return response

        except Exception as e:
            print(f"Error processing document: {str(e)}")
            return jsonify({'error': 'Lỗi khi xử lý tài liệu. Vui lòng thử lại.'}), 500

        finally:
            # Clean up temporary file
            try:
                if os.path.exists(temp_doc_path):
                    os.remove(temp_doc_path)
            except Exception as cleanup_error:
                print(f"Warning: Could not remove temporary file: {str(cleanup_error)}")

    except Exception as e:
        print(f"Error generating document: {str(e)}")
        return jsonify({'error': 'Có lỗi xảy ra khi tạo tài liệu. Vui lòng thử lại sau.'}), 500

def load_form_history():
    if os.path.exists(form_history_path):
        with open(form_history_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    return []

def save_form_history(data):
    with open(form_history_path, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=4)

@app.route('/form/<form_id>')
def view_form(form_id):
    try:
        # Load form history
        form_history = load_form_history()
        
        # Find the form with the matching ID
        form_data = None
        form_path = None
        for form in form_history:
            # Lấy ID từ đường dẫn file
            file_name = os.path.basename(form['path'])
            file_id = file_name.split('_')[0] if '_' in file_name else file_name.split('.')[0]
            
            # So sánh với form_id được truyền vào
            if file_id == form_id or file_name.startswith(form_id):
                form_data = form.get('form_data')
                form_path = form['path']
                break
        
        if not form_data or not form_path:
            return jsonify({'error': 'Form not found'}), 404
            
        # Set the global doc_path to the form's path
        global doc_path
        doc_path = form_path
        
        # Load the document and extract fields
        text = load_document(doc_path)
        fields = extract_fields(text)
        
        # Get suggestions
        db_data = load_db()
        suggestions = generate_suggestions(db_data) if db_data else {}
        
        # Render the form with pre-filled data
        return render_template("index.html", fields=fields, suggestions=suggestions, form_data=form_data)
        
    except Exception as e:
        print(f"Error loading form: {str(e)}")
        return jsonify({'error': 'Failed to load form'}), 500

@app.route('/delete-form/<form_id>', methods=['DELETE'])
def delete_form(form_id):
    try:
        # Load form history
        form_history = load_form_history()
        
        # Find the form with the matching ID
        form_index = None
        for i, form in enumerate(form_history):
            # Lấy ID từ đường dẫn file
            file_name = os.path.basename(form['path'])
            file_id = file_name.split('_')[0] if '_' in file_name else file_name.split('.')[0]
            
            # So sánh với form_id được truyền vào
            if file_id == form_id or file_name.startswith(form_id):
                form_index = i
                break
        
        if form_index is None:
            return jsonify({'error': 'Form not found'}), 404
        
        # Xóa form khỏi lịch sử
        deleted_form = form_history.pop(form_index)
        save_form_history(form_history)
        
        # Xóa file nếu cần
        try:
            if os.path.exists(deleted_form['path']) and os.path.isfile(deleted_form['path']):
                # Chỉ xóa file nếu nó nằm trong thư mục uploads
                if uploads_dir in deleted_form['path']:
                    os.remove(deleted_form['path'])
        except Exception as e:
            print(f"Warning: Could not delete file: {str(e)}")
        
        return jsonify({'message': 'Form deleted successfully'})
        
    except Exception as e:
        print(f"Error deleting form: {str(e)}")
        return jsonify({'error': 'Failed to delete form'}), 500

if __name__ == '__main__':
    app.run(debug=True)